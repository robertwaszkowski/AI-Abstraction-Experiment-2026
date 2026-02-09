
import os
import docx

def extract_text(file_path):
    print(f"--- START FILE: {os.path.basename(file_path)} ---")
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text.strip())
        
        # Also extract tables
        for table in doc.tables:
            print("[TABLE]")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip().replace('\n', ' '))
                print(" | ".join(row_data))
            print("[END TABLE]")

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    print(f"--- END FILE: {os.path.basename(file_path)} ---")

target_files = [
    "Change of Employment Conditions Data.docx"
]

for f in target_files:
    if os.path.exists(f):
        extract_text(f)
    else:
        print(f"File not found: {f}")
