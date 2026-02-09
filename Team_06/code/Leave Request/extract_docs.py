# -*- coding: utf-8 -*-
"""Script to extract contents from DOCX files"""

from docx import Document
import os

def extract_docx(filename, f):
    """Extract all text content from a DOCX file"""
    f.write(f"\n{'='*60}\n")
    f.write(f"FILE: {filename}\n")
    f.write('='*60 + "\n")
    
    doc = Document(filename)
    
    f.write("\n--- PARAGRAPHS ---\n")
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + "\n")
    
    f.write("\n--- TABLES ---\n")
    for idx, table in enumerate(doc.tables):
        f.write(f"\n=== Table {idx+1} ===\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

if __name__ == "__main__":
    with open("extracted_content.txt", "w", encoding="utf-8") as f:
        extract_docx("Leave Request Data.docx", f)
        extract_docx("Leave Request Test Scenario.docx", f)
    print("Content saved to extracted_content.txt")
