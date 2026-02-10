import docx

# Read Data document
print("=" * 60)
print("DECORATIONS AND MEDALS DATA")
print("=" * 60)
doc = docx.Document('Decorations and Medals Data.docx')
print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'{i}: {p.text}')
        
print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f'\nTable {t_idx}:')
    for r_idx, row in enumerate(table.rows):
        vals = [cell.text for cell in row.cells]
        print(f'  Row {r_idx}: {vals}')

# Read Test Scenario document
print("\n\n")
print("=" * 60)
print("DECORATIONS AND MEDALS TEST SCENARIO")
print("=" * 60)
doc2 = docx.Document('Decorations and Medals Test Scenario.docx')
print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip():
        print(f'{i}: {p.text}')
        
print("\n=== TABLES ===")
for t_idx, table in enumerate(doc2.tables):
    print(f'\nTable {t_idx}:')
    for r_idx, row in enumerate(table.rows):
        vals = [cell.text for cell in row.cells]
        print(f'  Row {r_idx}: {vals}')
